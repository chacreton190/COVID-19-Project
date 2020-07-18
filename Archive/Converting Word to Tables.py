import docx
import os
# from docx import document as doc


def ReadingTextDocuments(file_name):
    doc = docx.Document(file_name)

    completedText = []
    counter = 1
    for paragraph in doc.paragraphs:

        completedText.append(paragraph.text)
        print(f"Paragraph {counter}: {paragraph.text}")
        counter += 1
    return '\n' .join(completedText)

def ConvertWordtoTxt(inpath, outpath):
    doc = docx.Document(inpath)
    doc.save(outpath)

    for table in doc.tables:
        print(type(table))
        for row in table.rows:
            print(type(row))
            for col in row.cells:
                print(type(col))
                print(col.text)
filepath = "Data\Word\Pages from covid-19-data---daily-report-2020-03-16-1815.docx"
outfilepath ='Data\Txt\Text Version daily-report-2020-03-16.txt'
# ReadingTextDocuments(filepath)

ConvertWordtoTxt(filepath, outfilepath)
